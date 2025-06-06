import re
from docx import Document
from typing  import List, Dict
from pprint import pprint, pformat

from domain.llm_utils import LLMUtils
from domain.worker_class import Worker
from domain.queue import MetadataDoc
from infrastructure.open_microsoft_document import IOpenAndUpdateDocument
from infrastructure.generic_logger import GenericLogger
from docx.enum.style import WD_STYLE_TYPE

class OpenDOCDocument(IOpenAndUpdateDocument):
    HEADING_NAME: str = "Heading name"
    SECTION_TEXT: str = "Paragraph text"
    HEADING_POINTER: str = "Heading Pointer"
    PARAGRAPH_POINTERS: str = "Paragraph Pointers"
    INIT_HEADING = "no heading"
    def __init__(self, 
                 document_path: str,
                 worker: Worker, 
                 paragraph_start_min_word_numbers: int,
                 paragraph_start_min_word_length: int, 
                 logger: GenericLogger,  force_context: str = None):
        super().__init__(document_path, worker, 
                         paragraph_start_min_word_numbers, paragraph_start_min_word_length, 
                         logger)
        self.document =  Document(document_path)
        self.document_styles: List = [ s for s in self.document.styles if s.type in [WD_STYLE_TYPE.PARAGRAPH, WD_STYLE_TYPE.LIST] ]
        self.logger.log_info("Styles found in document:")
        self.force_context = force_context
        for style in self.document_styles:
            self.logger.log_info(f' * {style.name}')

    def __get_heading_deepness(self, heading_style) -> int:
        regexp = re.compile(r'^heading\s+(?P<heading_deepness>\d+)')
        m = regexp.match(heading_style)
        if m:
            return int(m.group('heading_deepness'))
        return -1

    def __prepend_heading_to_paragraph_text(self, heading_text, paragraph_text) -> str:
        return heading_text + "\n" + paragraph_text

    def __append_data(self, file_content: List, current_heading_style: str, highest_heading_style: str, \
                      latest_heading_pointer: any, heading_text: str, \
                      latest_paragraphs_pointer: List, paragraph_text: str):
        # When appending a new heading and assocaited text, we search for the latest element of file_content
        # The we iterate over all heading searching for latest element of each array until parent heading is found
        heading_deepness: int = self.__get_heading_deepness(current_heading_style)
        if heading_deepness > 0:
            heading_text = f"{'#' * heading_deepness} {heading_text}"
            self.logger.log_trace(f"Found heading {current_heading_style}, deepness: {heading_deepness}, text: {heading_text}")

        # paragraph_pointer: List = [latest_heading_pointer]
        # if latest_paragraphs_pointer is not None and len(latest_paragraphs_pointer) > 0:
        #     paragraph_pointer.extend(latest_paragraphs_pointer)
        new_data_structure: Dict = {
            current_heading_style:[
                {self.HEADING_NAME: heading_text, self.HEADING_POINTER: latest_heading_pointer},
                {self.SECTION_TEXT: self.__prepend_heading_to_paragraph_text(heading_text, paragraph_text), self.PARAGRAPH_POINTERS: latest_paragraphs_pointer}
            ]
        }
        self.logger.log_trace(f"new_data_structure: {pformat(new_data_structure, width=200)}")
        self.logger.log_trace(f"self.PARAGRAPH_POINTERS: {pformat(['Pointer: ' + str(pointer) + ', Text:' + pointer.text for pointer in new_data_structure[current_heading_style][1][self.PARAGRAPH_POINTERS]], width=200)}")

        if (len(file_content) > 0 or current_heading_style < highest_heading_style) and \
                            current_heading_style.startswith("heading"):
            # Search for latest data structure indexed by a header higher than current one
            sub_heading_found: bool = True
            parent_data_structure: List = file_content

            if len(parent_data_structure) > 0:
                while sub_heading_found:
                    sub_heading_found = False
                    for key, dictionary in parent_data_structure[-1].items():
                        self.logger.log_trace(f"In for loop: key: {key}, \ndictionnary: {str(pformat(dictionary, width=200))[:100]}\n\n")
                        if key.startswith("heading") and key < current_heading_style:
                            self.logger.log_trace(f"**** Heading {key} accepted****")
                            parent_data_structure = dictionary
                            # There can be only one key
                            sub_heading_found = True
                            break
            self.logger.log_trace(f"parent_data_structure: {pformat(parent_data_structure, width=200)}")
            parent_data_structure.append(new_data_structure)
        else:
            file_content.append(new_data_structure)
        self.logger.log_debug(f"file_content: {pformat(file_content, width=230)}")

    def __iter_headings(self, paragraphs):
        file_content: List = []
        latest_heading_style: str = self.INIT_HEADING
        highest_heading_style: str = self.INIT_HEADING
        latest_heading_text: str = self.INIT_HEADING
        latest_text_paragraph: str = ""
        latest_paragraphs_pointer: List = []
        latest_heading_pointer: any = None
        # Below is an example of the data structure that is going to be created
        # [{"Heading 1": [
        #     {"Heading name": "Heading ...", "Pointer": pointer}
        #     {"Text": "Full Text content below heading1", "Pointer": [pointer1, pointer2, ...]},
        #     {"Heading 2": [
        #         {"Heading name": "Heading ...", "Pointer": pointer}
        #         {"Text": "Full Text content below Heading 2", "Pointer": [pointer1, pointer2, ...]},
        #         {"Heading 3": [
        #           {"Heading name": "Heading ...", "Pointer": pointer}
        #           {"Text": "Full Text content below Heading 3", "Pointer": [pointer1, pointer2, ...]}
        #         ]},
        #         {"Heading 3": [
        #           {"Heading name": "Heading ...", "Pointer": pointer}
        #           {"Text": "Full Text content below Heading 3", "Pointer": [pointer1, pointer2, ...]},
        #          ]},
        #     ]}
        #  ]},
        #  {"Heading 1": [   // When appending search latest of each element until parent heading is found
        #     {"Heading name": "Heading ...", "Pointer": pointer}
        #     {"Text": "Full Text content below heading1", "Pointer": [pointer1, pointer2, ...]},
        #     {"Heading 2": [
        #         {"Heading name": "Heading ...", "Pointer": pointer}
        #         {"Text": "Full Text content below Heading 2", "Pointer": [pointer1, pointer2, ...]}]
        #     },
        #     {"Heading 2": [
        #         {"Heading name": "Heading ...", "Pointer": pointer}
        #         {"Text": "Full Text content below Heading 2", "Pointer": [pointer1, pointer2, ...]}
        #         {"Heading 3": [
        #           {"Heading name": "Heading ...", "Pointer": pointer}
        #           {"Text": "Full Text content below Heading 3", "Pointer": [pointer1, pointer2, ...]},
        #          ]}
        #       ]
        #     }
        #    ]
        #  }
        # ]

        for paragraph in paragraphs:
            current_heading_style: str = paragraph.style.name.lower()
            if current_heading_style.startswith('heading'):
                # We need to consider cases where a document starts with Heading 2 and later gets Heading 1
                if latest_heading_pointer is not None:
                    self.logger.log_trace(f"Adding data from latest_heading_text = {latest_heading_text}")
                    self.__append_data(file_content, latest_heading_style, highest_heading_style, \
                                       latest_heading_pointer, latest_heading_text, \
                                       latest_paragraphs_pointer, latest_text_paragraph)
                latest_text_paragraph = ""
                latest_paragraphs_pointer = []
               
                latest_heading_style = paragraph.style.name.lower()
                latest_heading_text = paragraph.text
                latest_heading_pointer = paragraph
                current_heading_deepness: int = self.__get_heading_deepness(current_heading_style)
                highest_heading_deepness: int = self.__get_heading_deepness(highest_heading_style)
                if highest_heading_style == self.INIT_HEADING or \
                   (current_heading_deepness < highest_heading_deepness and current_heading_deepness > 0):
                    highest_heading_style = current_heading_style
            else:
                # Typcally a document starts with a main title that we represent here as heading 0
                if latest_heading_pointer == None:
                    latest_heading_text = paragraph.text
                    latest_heading_pointer = paragraph
                    latest_heading_style = "heading 0"
                    self.logger.log_debug(f"Found paragraph before any heading: {latest_heading_text}, creating fakes heading: {latest_heading_style}")

                latest_text_paragraph += paragraph.text + "\n"
                latest_paragraphs_pointer.append(paragraph)
        self.__append_data(file_content, latest_heading_style, highest_heading_style, \
                           latest_heading_pointer, latest_heading_text, \
                           latest_paragraphs_pointer, latest_text_paragraph)
        self.logger.log_debug(f"File content returned: \n{pformat(file_content, width=230)}")
        return file_content
    
    def __get_context(self, prev_headings: List) -> str:
        if self.force_context is not None:
            return "\n".join(self.force_context)  
        else: 
            return "The context is a list of headings belonging to the document. This list is intended to provide guidance to the LLM:" +\
                      "\n".join(prev_headings)  
        
    def __dispatch_requests(self, file_content: List, prev_headings: List):
        paragraph_text: str = ""
        list_pointers: List = []
        heading_name: str = None
        request_type: str = LLMUtils.DEFAULT_REQUEST
        heading_found: bool = False
        next_prev_headings: List = []
        for key_in_section in [self.HEADING_POINTER, self.PARAGRAPH_POINTERS]:
            for section in file_content:
                if isinstance(section, dict):
                    if not key_in_section in section:
                        continue
                    if self.SECTION_TEXT in section and len(section[self.SECTION_TEXT]) > 0:
                        paragraph_text += section[self.SECTION_TEXT]
                    if self.HEADING_POINTER in section:
                        list_pointers.append(section[self.HEADING_POINTER])
                    if self.PARAGRAPH_POINTERS in section: 
                        list_pointers.extend(section[self.PARAGRAPH_POINTERS])

                    if self.HEADING_NAME in section:
                        heading_name = section[self.HEADING_NAME]
                        heading_found = True
                        next_prev_headings.append(heading_name)
                    
        if heading_found and paragraph_text == self.__prepend_heading_to_paragraph_text(heading_name, ""):
            request_type = LLMUtils.HEADING_REQUEST 

        context: str = self.__get_context(prev_headings)
        prev_headings.extend(next_prev_headings)
        self.logger.log_trace(f"Preparing text to be used for the request: {paragraph_text}")
        self.worker.add_work_element(MetadataDoc(list_pointers, \
                                                 context, \
                                                 paragraph_text,
                                                 request_type,
                                                 self.logger, self.document_styles))
        match = re.compile(r'^heading\s+(\d+)')
        for section in file_content:
            for key, sub_section in section.items(): 
                if match.match(key):
                    self.__dispatch_requests(sub_section, prev_headings)

    def __dispatch_requests_DELETEME(self, file_content: List, prev_headings: List):
        for section in file_content:
            heading_name: str = None
            if isinstance(section, dict):
                section_text: str = section[self.SECTION_TEXT] if self.SECTION_TEXT in section else ""
                list_pointers: List = [section[self.HEADING_POINTER]] if self.HEADING_POINTER in section else []
                list_pointers.extend(section[self.PARAGRAPH_POINTERS] if self.PARAGRAPH_POINTERS in section else [])
                request_type: str = LLMUtils.DEFAULT_REQUEST

                if self.HEADING_NAME in section:
                    heading_name: str = section[self.HEADING_NAME]
                    section_text = re.sub(r'^\s*$', '', section_text)
                    if len(section) == 0: 
                        request_type = LLMUtils.HEADING_REQUEST 
                    else:
                        section_text = heading_name + "\n" + section_text
                    prev_headings.append(heading_name)

                context: str = self.__get_context(prev_headings)
                self.logger.log_trace(f'Extracted text section for LLM request: {section_text}')
                self.worker.add_work_element(MetadataDoc(list_pointers, \
                                                         context, \
                                                         section_text,
                                                         request_type,
                                                         self.logger, self.document_styles))

                match = re.compile(r'^heading\s+(\d+)')
                for key, sub_section in section.items(): 
                    if match.match(key):
                        self.__dispatch_requests(sub_section, prev_headings)

    def __doc_table_to_md_table(self, doc_table):        
        md_table: str = ""
        first_row: bool = True
        for row in doc_table.rows:
            current_row: str = ""
            for cell in row.cells:
                self.logger.log_debug("Transforming DOC table to MD Table")
                cell_value: str = ""
                if len(cell.paragraphs) > 0:
                    cell_value = " ".join([paragraph_pointer.text for paragraph_pointer in cell.paragraphs]) + "\n"
                    for paragraph_pointer in cell.paragraphs:
                        paragraph_pointer.text = ""
                    
                if len(cell.tables) > 0:
                    for doc_table in cell.tables:
                        md_table_in_cell_value: str = self.__doc_table_to_md_table(cell) + "\n"
                        cell_value += md_table_in_cell_value

                current_row += f"|{cell_value}"
            if len(current_row) > 0:
                md_table += "\n" + current_row + "|"
            if first_row:
                md_table = "\n" + "-" * len(md_table)
                first_row = False
        return md_table

    #TODO: All requests should be running in multiple threads
    def __fill_tasks(self, document: any):
        file_content: List = self.__iter_headings(document.paragraphs)
        prev_headings: List = []
        self.__dispatch_requests(file_content, prev_headings)

        for doc_table in document.tables:
            md_table = self.__doc_table_to_md_table(doc_table)
            context: str = self.__get_context(prev_headings)

            # In word dpcument we are going to replace a table 
            self.worker.add_work_element(MetadataDoc([doc_table], \
                                                      context, \
                                                      md_table,
                                                      LLMUtils.TABLE_REQUEST,
                                                      self.logger, self.document_styles
                                                      ))            

    def process(self):
        self.__fill_tasks(self.document)
        self.worker.process_all()

